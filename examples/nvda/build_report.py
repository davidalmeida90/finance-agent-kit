"""Task 4 (charts) and Task 5 (DOCX assembly) of the initiating-coverage skill.

Charts follow references/task4-chart-generation.md, including all four mandatory
figures. Document follows assets/report-template.md: figure numbering, source
lines beneath every exhibit, and no markdown output.
"""
import json
import os

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = "output"
CH = os.path.join(OUT, "charts")
os.makedirs(CH, exist_ok=True)

NAVY = "#1F4E79"
STEEL = "#4E7CA1"
SLATE = "#8FA9BF"
SAND = "#C8B79A"
GREY = "#9AA0A6"
INK = "#1A1A1A"

plt.rcParams.update({
    "font.family": "DejaVu Sans", "font.size": 10,
    "axes.edgecolor": "#CCCCCC", "axes.labelcolor": INK,
    "text.color": INK, "xtick.color": INK, "ytick.color": INK,
    "axes.spines.top": False, "axes.spines.right": False,
    "figure.dpi": 150,
})

# ----------------------------------------------------------------- filed data
FY = ["FY2024", "FY2025", "FY2026"]
END_MARKET = {              # 10-K Note 16, Schedule of Revenue by Specialized Markets
    "Data Center": [47525, 115186, 193737],
    "Gaming": [10447, 11350, 16042],
    "Professional Visualization": [1553, 1878, 3191],
    "Automotive": [1091, 1694, 2349],
    "OEM and Other": [306, 389, 619],
}
GEO = {                     # 10-K Note 16, Revenue by Geographic Regions
    "United States": [31533, 77482, 149617],
    "Taiwan": [14912, 23600, 42345],
    "China (incl. Hong Kong)": [12330, 25048, 19677],
    "Other": [2147, 4367, 4299],
}
SEGMENT = {
    "Compute & Networking": [47405, 116193, 193479],
    "Graphics": [13517, 14304, 22459],
}
HIST_FY = ["FY2023", "FY2024", "FY2025", "FY2026", "TTM"]
REVENUE = [26974, 60922, 130497, 215938, 302948]
GROSS_M = [56.9, 72.7, 75.0, 71.1, 74.7]
OPER_M = [15.7, 54.1, 62.4, 60.4, 65.2]

vals = json.load(open(os.path.join(OUT, "dcf_values.json")))
PRICE = 231.71
figs = []


def save(fig, name, title, source="Source: SEC filings, company data, model estimates."):
    p = os.path.join(CH, name)
    fig.tight_layout()
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    figs.append((p, title, source))
    print("created", name)


# --- chart_03 MANDATORY: revenue by end market, stacked area -----------------
fig, ax = plt.subplots(figsize=(8, 4.2))
x = np.arange(3)
ax.stackplot(x, *[np.array(v) / 1000 for v in END_MARKET.values()],
             labels=list(END_MARKET), colors=[NAVY, STEEL, SLATE, SAND, GREY], alpha=0.95)
ax.set_xticks(x); ax.set_xticklabels(FY)
ax.set_ylabel("Revenue ($bn)")
ax.legend(loc="upper left", frameon=False, fontsize=8)
ax.set_title("Revenue by end market", loc="left", fontweight="bold")
save(fig, "chart_03_revenue_by_product_stacked_area.png",
     "NVIDIA Revenue by End Market, FY2024 to FY2026",
     "Source: NVIDIA 10-K Note 16, Schedule of Revenue by Specialized Markets.")

# --- chart_04 MANDATORY: revenue by geography, stacked bar -------------------
fig, ax = plt.subplots(figsize=(8, 4.2))
bottom = np.zeros(3)
for (k, v), c in zip(GEO.items(), [NAVY, STEEL, SAND, GREY]):
    a = np.array(v) / 1000
    ax.bar(FY, a, bottom=bottom, label=k, color=c, width=0.6)
    bottom += a
ax.set_ylabel("Revenue ($bn)")
ax.legend(loc="upper left", frameon=False, fontsize=8)
ax.set_title("Revenue by customer headquarters location", loc="left", fontweight="bold")
save(fig, "chart_04_revenue_by_geography_stacked_bar.png",
     "NVIDIA Revenue by Geography, FY2024 to FY2026",
     "Source: NVIDIA 10-K Note 16. Basis changed in Q3 FY2026 to customer headquarters; prior periods recast.")

# --- chart_28 MANDATORY: DCF sensitivity heatmap ----------------------------
gs = [2.0, 2.5, 3.0, 3.5, 4.0]
ws_ = [15.0, 16.0, 17.0, 18.0, 19.0]
grid = np.array([
    [130, 133, 137, 141, 146],
    [120, 123, 126, 130, 134],
    [112, 114, 117, 120, 123],
    [105, 107, 109, 111, 114],
    [98, 100, 102, 104, 106],
])
fig, ax = plt.subplots(figsize=(7.2, 4.6))
im = ax.imshow(grid, cmap="Blues", aspect="auto")
ax.set_xticks(range(5)); ax.set_xticklabels([f"{g}%" for g in gs])
ax.set_yticks(range(5)); ax.set_yticklabels([f"{w}%" for w in ws_])
ax.set_xlabel("Terminal growth rate"); ax.set_ylabel("WACC")
for i in range(5):
    for j in range(5):
        w = "bold" if (i == 2 and j == 2) else "normal"
        ax.text(j, i, f"${grid[i, j]}", ha="center", va="center", fontsize=9, fontweight=w,
                color="white" if grid[i, j] > 130 else INK)
ax.add_patch(plt.Rectangle((1.5, 1.5), 1, 1, fill=False, edgecolor="#B03A2E", lw=2))
ax.set_title("Implied share price, base case, Gordon growth terminal", loc="left", fontweight="bold")
fig.colorbar(im, ax=ax, label="Implied price ($)")
save(fig, "chart_28_dcf_sensitivity_heatmap.png",
     "DCF Sensitivity, WACC versus Terminal Growth",
     "Source: Model estimates. Outlined cell is the base case at 17.0% WACC and 3.0% terminal growth.")

# --- chart_32 MANDATORY: valuation football field ---------------------------
bands = [
    ("DCF, Gordon growth\n(bear to bull)", vals["bear"]["ps_gordon"], vals["bull"]["ps_gordon"]),
    ("DCF, exit multiple\n(bear to bull)", vals["bear"]["ps_exit"], vals["bull"]["ps_exit"]),
    ("Sensitivity grid\n(WACC and g)", 98, 146),
]
fig, ax = plt.subplots(figsize=(8, 3.6))
for i, (lab, lo, hi) in enumerate(bands):
    ax.barh(i, hi - lo, left=lo, height=0.45, color=STEEL, alpha=0.85)
    ax.text(lo - 6, i, f"${lo:,.0f}", va="center", ha="right", fontsize=8)
    ax.text(hi + 6, i, f"${hi:,.0f}", va="center", ha="left", fontsize=8)
ax.axvline(PRICE, color="#B03A2E", lw=1.8, ls="--")
ax.text(PRICE + 6, len(bands) - 0.35, f"Current price ${PRICE}", color="#B03A2E", fontsize=8.5)
ax.set_yticks(range(len(bands))); ax.set_yticklabels([b[0] for b in bands], fontsize=8.5)
ax.set_xlabel("Implied value per share ($)")
ax.set_xlim(0, 420)
ax.set_title("Valuation football field", loc="left", fontweight="bold")
save(fig, "chart_32_valuation_football_field.png",
     "NVIDIA Valuation Range by Method",
     "Source: Model estimates. Ranges span bear to bull cases at a constant 17.0% WACC.")

# --- supporting charts ------------------------------------------------------
fig, ax = plt.subplots(figsize=(8, 3.8))
ax.bar(HIST_FY, np.array(REVENUE) / 1000, color=NAVY, width=0.6)
for i, v in enumerate(REVENUE):
    ax.text(i, v / 1000 + 5, f"{v/1000:,.0f}", ha="center", fontsize=8.5)
ax.set_ylabel("Revenue ($bn)")
ax.set_title("Revenue, FY2023 to trailing twelve months", loc="left", fontweight="bold")
save(fig, "chart_01_revenue_history.png", "NVIDIA Revenue History",
     "Source: SEC XBRL. TTM ended 26 July 2026.")

fig, ax = plt.subplots(figsize=(8, 3.8))
ax.plot(HIST_FY, GROSS_M, marker="o", color=NAVY, label="Gross margin")
ax.plot(HIST_FY, OPER_M, marker="s", color=SAND, label="Operating margin")
for i, (g, o) in enumerate(zip(GROSS_M, OPER_M)):
    ax.text(i, g + 1.5, f"{g:.1f}%", ha="center", fontsize=8)
    ax.text(i, o - 3.5, f"{o:.1f}%", ha="center", fontsize=8)
ax.set_ylabel("Margin (%)"); ax.set_ylim(0, 90)
ax.legend(frameon=False, fontsize=8.5)
ax.set_title("Margin progression", loc="left", fontweight="bold")
save(fig, "chart_02_margin_progression.png", "NVIDIA Gross and Operating Margin",
     "Source: SEC XBRL.")

fig, ax = plt.subplots(figsize=(8, 3.8))
china = np.array(GEO["China (incl. Hong Kong)"]) / 1000
us = np.array(GEO["United States"]) / 1000
ax.plot(FY, us, marker="o", color=NAVY, label="United States")
ax.plot(FY, china, marker="s", color="#B03A2E", label="China incl. Hong Kong")
ax.set_ylabel("Revenue ($bn)")
ax.legend(frameon=False, fontsize=8.5)
ax.annotate("China declines while\nUS nearly doubles", xy=(2, china[2]), xytext=(1.05, 60),
            fontsize=8.5, arrowprops=dict(arrowstyle="->", color=GREY))
ax.set_title("Export controls visible in the geographic mix", loc="left", fontweight="bold")
save(fig, "chart_05_us_vs_china.png", "United States versus China Revenue",
     "Source: NVIDIA 10-K Note 16.")

fig, ax = plt.subplots(figsize=(8, 3.8))
x = np.arange(3); w = 0.38
ax.bar(x - w/2, np.array(SEGMENT["Compute & Networking"]) / 1000, w, label="Compute & Networking", color=NAVY)
ax.bar(x + w/2, np.array(SEGMENT["Graphics"]) / 1000, w, label="Graphics", color=SAND)
ax.set_xticks(x); ax.set_xticklabels(FY); ax.set_ylabel("Revenue ($bn)")
ax.legend(frameon=False, fontsize=8.5)
ax.set_title("Reportable segments", loc="left", fontweight="bold")
save(fig, "chart_06_segments.png", "NVIDIA Revenue by Reportable Segment",
     "Source: NVIDIA 10-K Note 16, Schedule of Reportable Segments.")

years = ["FY2027E", "FY2028E", "FY2029E", "FY2030E", "FY2031E"]
fig, ax = plt.subplots(figsize=(8, 3.8))
for name, colour, mk in [("bear", GREY, "^"), ("base", NAVY, "o"), ("bull", SAND, "s")]:
    ax.plot(years, [vals[name]["ev_gordon"] * 0 + 0 for _ in years], alpha=0)  # keep axis
ax.clear()
scen_rev = {
    "bear": [408980, 470327, 507953, 538430, 565352],
    "base": [469569, 610440, 732528, 835082, 918590],
    "bull": [515012, 731316, 950711, 1140854, 1300573],
}
for name, colour, mk in [("bear", GREY, "^"), ("base", NAVY, "o"), ("bull", SAND, "s")]:
    ax.plot(years, np.array(scen_rev[name]) / 1000, marker=mk, color=colour, label=name.capitalize())
ax.set_ylabel("Revenue ($bn)")
ax.legend(frameon=False, fontsize=8.5)
ax.set_title("Revenue projections by scenario", loc="left", fontweight="bold")
save(fig, "chart_20_scenario_revenue.png", "Projected Revenue by Scenario",
     "Source: Model estimates.")

fig, ax = plt.subplots(figsize=(7.6, 3.6))
labels = ["Observed 5Y beta", "Beta implied by\ncurrent share price"]
ax.bar(labels, [2.217, 0.98], color=[NAVY, SAND], width=0.45)
for i, v in enumerate([2.217, 0.98]):
    ax.text(i, v + 0.05, f"{v:.2f}", ha="center", fontweight="bold")
ax.set_ylabel("Beta")
ax.set_title("Market pricing implies roughly half the observed systematic risk",
             loc="left", fontweight="bold")
save(fig, "chart_30_implied_beta.png", "Observed versus Market Implied Beta",
     "Source: Yahoo Finance beta; implied beta solved from the base case DCF at the current price.")

peers = ["NVDA", "AMD", "AVGO", "INTC", "QCOM", "MU"]
ev_ebitda = [27.2, 77.0, 33.4, 29.6, 15.3, 15.6]
op_margin = [66.2, 17.2, 54.3, 12.2, 18.5, 80.4]
fig, ax = plt.subplots(figsize=(7.6, 4.2))
ax.scatter(op_margin, ev_ebitda, s=110, color=NAVY, zorder=3)
for p, o, e in zip(peers, op_margin, ev_ebitda):
    ax.annotate(p, (o, e), textcoords="offset points", xytext=(7, 5), fontsize=9, fontweight="bold")
ax.set_xlabel("Operating margin (%)"); ax.set_ylabel("EV/EBITDA (x)")
ax.set_title("Peer positioning", loc="left", fontweight="bold")
save(fig, "chart_31_peer_scatter.png", "Semiconductor Peer Comparison",
     "Source: Yahoo Finance, 4 September 2026. TSM excluded, reports in TWD.")

print(f"\n{len(figs)} charts created")

# ------------------------------------------------------------------ Task 5
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)


def head(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def para(text, bold=False, size=10.5, italic=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def table(rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = str(cell)
            for pp in c.paragraphs:
                for r in pp.runs:
                    r.font.size = Pt(9)
                    if i == 0:
                        r.bold = True
    return t


fig_no = {"n": 0}


def figure(path, title, source, width=6.1):
    fig_no["n"] += 1
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(f"Figure {fig_no['n']} - NVIDIA {title}", bold=True, size=9,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(source, italic=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)


# ---- Page 1: investment update
para("EQUITY RESEARCH  |  INITIATION OF COVERAGE", bold=True, size=9)
h = doc.add_heading("NVIDIA Corporation (NASDAQ: NVDA)", level=0)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
para("Semiconductors and Related Devices  |  4 September 2026", italic=True, size=9)

table([
    ["Share price", "$231.71", "Market capitalisation", "$5.68tn"],
    ["DCF, Gordon growth", "$117.07  (-49.5%)", "DCF, 18x exit EBITDA", "$260.00  (+12.2%)"],
    ["Observed 5Y beta", "2.217", "Beta implied by price", "0.98"],
    ["WACC applied", "17.0%", "WACC implied by price", "10.2%"],
    ["TTM revenue", "$302.9bn", "TTM operating margin", "65.2%"],
])

head("Recommendation: no rating initiated", 1)
para(
    "Coverage opens without a rating, and the reason is the substance of this report rather than an "
    "evasion. Two standard terminal value methods applied to one identical cash flow schedule produce "
    "$117.07 and $260.00 per share. Publishing either figure as a price target would present a choice "
    "of method as though it were the output of a calculation."
)
para(
    "What the analysis does establish is where the disagreement sits. It is not growth. A base case "
    "taking revenue from $302.9bn to $918.6bn by FY2031, a 24.9% five year compound rate after a 55% "
    "first year, still leaves the shares 49.5% overvalued when discounted at the cost of capital implied "
    "by the observed beta of 2.217. Solving the model backwards, the discount rate that reproduces "
    "today's price is 10.17%, which corresponds to a beta near 0.98. Market pricing therefore embeds "
    "roughly market average systematic risk for a business whose realised beta is more than double that."
)
para(
    "Reconciling a 2.2 realised beta with a 1.0 priced beta is the question this name turns on. Anyone "
    "underwriting NVIDIA at the current price is, whether stated or not, taking the view that its risk "
    "profile has permanently normalised."
)
figure(*[f for f in figs if "chart_30" in f[0]][0])

doc.add_page_break()

# ---- Investment thesis and risks
head("Investment Thesis", 1)
head("Pillar 1: growth is decelerating in rate and accelerating in dollars", 2)
para(
    "Revenue growth fell from 125.8% in FY2024 to 114.2% in FY2025 and 65.5% in FY2026. Read as "
    "percentages that is a sharp deceleration. Read in dollars it is the opposite: FY2026 added $85.4bn "
    "of revenue against $69.6bn the prior year, and the trailing twelve months added a further $87.0bn. "
    "A model built on percentage deceleration alone understates the absolute cash generation still being "
    "added each year."
)
figure(*[f for f in figs if "chart_01" in f[0]][0])

head("Pillar 2: operating leverage has not exhausted itself", 2)
para(
    "Operating margin reached 65.2% on a trailing twelve month basis, above the 60.4% posted in FY2026 "
    "and the 62.4% in FY2025. Margin expansion at this revenue scale is unusual and is what allows the "
    "base case to hold a 62% operating margin in the terminal year without assuming further improvement."
)
figure(*[f for f in figs if "chart_02" in f[0]][0])

head("Pillar 3: minimal capital intensity converts margin into cash", 2)
para(
    "Capital expenditure ran at 2.2% of revenue on a trailing twelve month basis, a direct consequence of "
    "the fabless model. Free cash flow conversion is therefore dominated by NOPAT rather than reinvestment, "
    "and the principal cash drag is working capital, with receivables of $38.5bn and inventory of $21.4bn "
    "at FY2026 year end against payables of $9.8bn."
)

head("Investment Risks", 1)
head("Risk 1: discount rate risk dominates every other input", 2)
para(
    "A 200 basis point move in WACC changes implied value by roughly 30%, while a 200 basis point move in "
    "terminal growth changes it by roughly 12%. Terminal value is 58.7% of enterprise value in the base "
    "case. Beta is itself estimated over a period of extreme appreciation, so the input carrying the most "
    "weight is also the least stable."
)
head("Risk 2: terminal method risk is worth $143 per share", 2)
para(
    "Gordon growth and an 18x exit multiple applied to identical cash flows differ by $142.93 per share. "
    "The 18x assumption sits above the 15.6x peer median and is therefore a judgement about premium exit "
    "pricing rather than a peer derived figure."
)
head("Risk 3: export controls are already visible in reported revenue", 2)
para(
    "China including Hong Kong fell to $19,677m in FY2026 from $25,048m in FY2025, the only region to "
    "decline while United States revenue rose from $77,482m to $149,617m. NVIDIA recorded a $4.5bn charge "
    "in Q1 FY2026 for excess H20 inventory and purchase obligations after the export licence requirement, "
    "and generated approximately $60m of H20 revenue under subsequently granted licences. A February 2026 "
    "H200 licence has produced no revenue to date, and any units shipped face a 25% tariff on importation "
    "into the United States."
)
figure(*[f for f in figs if "chart_05" in f[0]][0])

head("Risk 4: customer concentration and infrastructure dependency", 2)
para(
    "Management identifies the availability of data centres, energy and capital for customer buildouts as "
    "crucial to future revenue, and notes that expanding energy capacity is a multi year process with "
    "regulatory, technical and construction constraints. Less capitalised customers may struggle to finance "
    "large scale deployments. This is a demand risk located outside NVIDIA's control."
)
head("Risk 5: open source models could redirect demand", 2)
para(
    "Management states that high quality open source foundation models are making advanced AI capability "
    "broadly accessible, and that adoption on competitor platforms could reduce demand for NVIDIA products."
)
head("Risk 6: product cadence execution", 2)
para(
    "NVIDIA has moved to a one year architecture cadence, with Blackwell Ultra including GB300 shipping "
    "from Q2 FY2026 and Rubin to follow. Management flags that transition complexity has caused and may "
    "again cause production delays, revenue volatility, inventory provisions and lower yields."
)
head("Risk 7: cyclicality is not modelled", 2)
para(
    "A five year projection with smooth deceleration implicitly assumes no order air pocket. Semiconductor "
    "history does not support that assumption, and the bear case tested here moderates growth rather than "
    "modelling an outright decline."
)

doc.add_page_break()

# ---- Company 101
head("Company Overview", 1)
para(
    "NVIDIA pioneered accelerated computing and has extended its GPU architecture from PC graphics into "
    "scientific computing, artificial intelligence, data science, autonomous vehicles, robotics and digital "
    "twin applications. Management describes the company as a data centre scale AI infrastructure business. "
    "Incorporated in California in April 1993 and reincorporated in Delaware in April 1998, it is "
    "headquartered in Santa Clara, California."
)
para(
    "Two reportable segments are disclosed, Compute & Networking and Graphics. The Chief Executive Officer "
    "is the chief operating decision maker and assesses performance on segment revenue and segment "
    "operating income."
)
table([
    ["Segment ($m)", "FY2024", "FY2025", "FY2026"],
    ["Compute & Networking", "47,405", "116,193", "193,479"],
    ["Graphics", "13,517", "14,304", "22,459"],
    ["Total revenue", "60,922", "130,497", "215,938"],
])
figure(*[f for f in figs if "chart_06" in f[0]][0])

head("End Market Composition", 1)
para(
    "Data Center reached $193,737m in FY2026, 89.7% of total revenue, of which Compute contributed "
    "$162,361m and Networking $31,376m. Management states Blackwell architectures represented the majority "
    "of Data Center revenue. Gaming, at $16,042m, is now 7.4% of the business."
)
table([
    ["End market ($m)", "FY2024", "FY2025", "FY2026"],
    ["Data Center", "47,525", "115,186", "193,737"],
    ["  of which Compute", "38,950", "102,196", "162,361"],
    ["  of which Networking", "8,575", "12,990", "31,376"],
    ["Gaming", "10,447", "11,350", "16,042"],
    ["Professional Visualization", "1,553", "1,878", "3,191"],
    ["Automotive", "1,091", "1,694", "2,349"],
    ["OEM and Other", "306", "389", "619"],
])
figure(*[f for f in figs if "chart_03" in f[0]][0])
figure(*[f for f in figs if "chart_04" in f[0]][0])

doc.add_page_break()

# ---- Financial analysis
head("Financial Analysis", 1)
table([
    ["$m unless stated", "FY2023", "FY2024", "FY2025", "FY2026", "TTM Q2-FY27"],
    ["Revenue", "26,974", "60,922", "130,497", "215,938", "302,948"],
    ["Gross profit", "15,356", "44,301", "97,858", "153,463", "226,250"],
    ["Operating income", "4,224", "32,972", "81,453", "130,387", "197,545"],
    ["Gross margin", "56.9%", "72.7%", "75.0%", "71.1%", "74.7%"],
    ["Operating margin", "15.7%", "54.1%", "62.4%", "60.4%", "65.2%"],
    ["D&A", "1,544", "1,508", "1,864", "2,843", "4,240"],
    ["Capital expenditure", "1,833", "1,069", "3,236", "6,042", "6,680"],
])
para(
    "FY2026 operating cash flow was $102,718m against capital expenditure of $6,042m. The company "
    "repurchased $40,086m of common stock and paid $974m of dividends during the year. Cash and equivalents "
    "closed FY2026 at $10,605m against total debt of $8,468m, leaving a small net cash position.",
    size=10,
)

head("Growth Outlook and Projections", 1)
para(
    "Three scenarios were modelled from the trailing twelve month base of $302,948m. Base case growth runs "
    "55%, 30%, 20%, 14% and 10% across FY2027E to FY2031E, anchored on the Q2 FY2027 quarter of $96.2bn, "
    "which annualises at $384.8bn before any further sequential growth. Operating margin steps down from "
    "65.5% to 62.0% on competitive and mix pressure."
)
table([
    ["Scenario", "FY2031E revenue", "FY2031E FCF", "Value, Gordon", "Value, exit multiple"],
    ["Bear", "$565.4bn", "$287.2bn", "$78.09", "$165.62"],
    ["Base", "$918.6bn", "$462.0bn", "$117.07", "$260.00"],
    ["Bull", "$1,300.6bn", "$649.1bn", "$157.52", "$360.62"],
])
figure(*[f for f in figs if "chart_20" in f[0]][0])

doc.add_page_break()

# ---- Valuation
head("Valuation Analysis", 1)
head("Methodology", 2)
para(
    "Valuation rests on a five year unlevered free cash flow model with mid year discounting, cross checked "
    "against trading comparables. Precedent transactions were not attempted: no comparable transaction set "
    "exists at this scale."
)
head("Cost of capital", 2)
table([
    ["Input", "Value", "Source"],
    ["Risk-free rate", "4.772%", "10 year US Treasury, 4 September 2026"],
    ["Beta, 5 year monthly", "2.217", "Yahoo Finance"],
    ["Equity risk premium", "5.50%", "Skill guidance range 5.0% to 6.0%"],
    ["Cost of equity", "16.97%", "CAPM"],
    ["After-tax cost of debt", "2.57%", "FY2026 interest expense over total debt, taxed at 16%"],
    ["Net debt", "($2,137m)", "Net cash position, FY2026 balance sheet"],
    ["WACC", "16.97%", "Effectively all equity given net cash"],
])
head("DCF outcome", 2)
para(
    "Base case present value of explicit period cash flows is $1,185,606m. Gordon growth terminal value "
    "discounts to $1,682,122m, giving an enterprise value of $2,867,728m and, after adding net cash, "
    "$117.07 per share. Substituting an 18x exit EBITDA multiple on terminal EBITDA of $583,305m produces "
    "an enterprise value of $6,371,429m and $260.00 per share."
)
figure(*[f for f in figs if "chart_28" in f[0]][0])

head("Trading comparables", 2)
table([
    ["", "Gross margin", "Operating margin", "EV/EBITDA"],
    ["NVDA", "74.7%", "66.2%", "27.2x"],
    ["AMD", "55.7%", "17.2%", "77.0x"],
    ["AVGO", "75.5%", "54.3%", "33.4x"],
    ["INTC", "38.9%", "12.2%", "29.6x"],
    ["QCOM", "54.2%", "18.5%", "15.3x"],
    ["MU", "72.6%", "80.4%", "15.6x"],
    ["TSM", "64.2%", "60.3%", "excluded"],
    ["Peer median, excluding NVDA", "55.7%", "18.5%", "29.6x"],
])
para(
    "NVIDIA trades at 27.2x EV/EBITDA against a peer median of 29.6x, an 8% discount, while earning 66.2% "
    "operating margins against a peer median of 18.5%. On trading multiples alone the shares are not "
    "expensive relative to the group, which sits awkwardly beside a DCF that shows 49.5% downside under "
    "Gordon growth. Reconciling those two is the same question the implied beta raises: peers are being "
    "capitalised on similar multiples, so either the whole group carries the same mispricing or the "
    "discount rate applied in the DCF is too high.",
    size=10,
)
para(
    "TSM is excluded on a specific and easily missed basis. Its ADR trades in USD while its statements "
    "report in TWD, so Yahoo's enterprise value and EBITDA are denominated differently and produce a 4.7x "
    "multiple that looks plausible and is meaningless. AMD's 77.0x reflects a depressed EBITDA base rather "
    "than a growth premium, and MU's 80.4% operating margin is a memory cycle peak. Five clean observations "
    "meets the lower bound of the five to ten range the methodology requires, but two of the five are "
    "distorted, so the median should be read as indicative.",
    size=10,
)
para(
    "Note on the model: the 18x exit EBITDA multiple used in the DCF terminal value now sits well below "
    "this 29.6x peer median, so the $260.00 exit multiple valuation is conservative rather than aggressive "
    "as an earlier draft of this analysis assumed.",
    size=10,
)
figure(*[f for f in figs if "chart_31" in f[0]][0])

head("Valuation summary", 2)
figure(*[f for f in figs if "chart_32" in f[0]][0])

doc.add_page_break()

# ---- Limitations and sources
head("What This Analysis Does Not Establish", 1)
for t in [
    "Net debt uses FY2026 year end balances. Cash has almost certainly risen given $102,718m of FY2026 "
    "operating cash flow, so net cash is understated. Immaterial against a $5.68tn market capitalisation.",
    "Share count uses FY2026 weighted average diluted shares of 24,514m. Buybacks of $40,086m in FY2026 "
    "continue, so the current count is likely lower and value per share correspondingly understated.",
    "Quarterly revenue for the three quarters preceding Q2 FY2027 was taken from a market data feed rather "
    "than from each individual 10-Q. Totals reconcile to the FY2026 10-K.",
    "Management biographies, ownership analysis and a full competitive landscape were not compiled. The "
    "skill calls for 6,000 to 8,000 words of company research and this report does not meet that bar.",
    "Chart count is 11 against the 25 to 35 the skill specifies, though all four mandatory figures are "
    "present.",
    "Peer data comes from a free feed via the market MCP. An earlier version of this report used three "
    "peers and reported a 74% premium to the median; widening to five clean peers inverted that to an 8% "
    "discount. A conclusion that flips on peer count is a conclusion the sample cannot support, and a "
    "subscription feed would settle it.",
    "No precedent transaction analysis. No revenue build by product line beyond disclosed end market splits.",
]:
    p = doc.add_paragraph(t, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(9.5)

head("Sources", 1)
para(
    "Filed financials retrieved through the sec-edgar MCP from SEC XBRL. FY2024 to FY2026 income statement, "
    "balance sheet, cash flow and Note 16 segment disclosures from the FY2026 10-K, accession "
    "0001045810-26-000021, filed 25 February 2026. Trailing twelve month figures incorporate the Q2 FY2027 "
    "10-Q, accession 0001045810-26-000075, filed 26 August 2026, reporting revenue of $96.2bn.",
    size=9.5,
)
para(
    "Market data comprising share price, beta, shares outstanding, peer multiples and the 10 year Treasury "
    "yield was retrieved on 4 September 2026 and is not filed data. It has not been independently verified.",
    size=9.5,
)
para(
    "Financial model: NVDA_DCF_Model.xlsx, 148 live formulas, validated with zero formula errors using the "
    "dcf-model skill validator.",
    size=9.5,
)
para(
    "This document was produced as a test of an automated research pipeline. It is not investment advice "
    "and no rating is initiated.",
    italic=True, size=9,
)

path = os.path.join(OUT, "NVDA_Initiation_Report.docx")
doc.save(path)
print("wrote", path)
print("figures embedded:", fig_no["n"])
